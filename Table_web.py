import streamlit as st
from PIL import Image, ImageDraw, ImageFont
import textwrap
import io
import os
import math
import zipfile
import re
from datetime import date

# Try to set dark theme for Streamlit
try:
    cfg_dir = os.path.join(os.getcwd(), ".streamlit")
    os.makedirs(cfg_dir, exist_ok=True)
    cfg_path = os.path.join(cfg_dir, "config.toml")
    with open(cfg_path, "w", encoding="utf-8") as f:
        f.write('[theme]\nbase = "dark"\n')
except Exception:
    pass

# ---------- Shared helpers ----------
def safe_filename(s):
    if s is None:
        return ""
    s = str(s).strip()
    s = s.replace(" ", "_")
    return re.sub(r"[^A-Za-z0-9._-]", "", s)

def wrap_text(text, width):
    if text is None or str(text).strip() == "":
        return {"wrapped_text": "", "line_count": 0}
    words = str(text).split(" ")
    wrapped_lines = []
    current_line = ""
    lines_for_curr_text = 0
    for word in words:
        if current_line == "":
            projected_len = len(word)
        else:
            projected_len = len(current_line) + 1 + len(word)
        if projected_len > width:
            wrapped_lines.append(current_line)
            current_line = word
            lines_for_curr_text += 1
        else:
            current_line = word if current_line == "" else current_line + " " + word
    if current_line != "":
        wrapped_lines.append(current_line)
        lines_for_curr_text += 1
    return {"wrapped_text": "\n".join(wrapped_lines), "line_count": lines_for_curr_text}

# ---------- TableImage (for the "Tabel Maken" preview) ----------
class TableImage:
    def __init__(
        self,
        rows,
        cols,
        row_height,
        col_width,
        font_path="verdana.ttf",
        bold_font_path="verdanab.ttf",
        font_size=10,
        line_color=(0, 0, 0),
        bg_color=(255, 255, 255),
        line_width=2,
        wrap_width=30,
        padding_left=8,
    ):
        self.rows = int(rows)
        self.cols = int(cols)
        self.col_width = int(col_width)
        self.line_color = line_color
        self.bg_color = bg_color
        self.line_width = int(line_width)
        self.wrap_width = wrap_width
        self.padding_left = int(padding_left)
        if isinstance(row_height, list):
            if len(row_height) != self.rows:
                if len(row_height) < self.rows:
                    row_height = row_height + [row_height[-1]] * (self.rows - len(row_height))
                else:
                    row_height = row_height[: self.rows]
            self.row_height = [int(h) for h in row_height]
        else:
            self.row_height = [int(row_height)] * self.rows
        self.height = sum(self.row_height)
        self.width = int(self.cols * self.col_width)
        try:
            BASE_DIR = os.path.dirname(os.path.abspath(__file__))
        except Exception:
            BASE_DIR = os.getcwd()
        self.font_path = os.path.join(BASE_DIR, font_path)
        self.bold_font_path = os.path.join(BASE_DIR, bold_font_path)
        self.font_size = int(font_size)
        self.cells = {}
        self.bold_cells = {}
    def set_text(self, row, col, text, bold=False):
        self.cells[(int(row), int(col))] = "" if text is None else str(text)
        self.bold_cells[(int(row), int(col))] = bool(bold)
    def draw(self):
        img = Image.new("RGB", (self.width, self.height), self.bg_color)
        draw = ImageDraw.Draw(img)
        y_positions = [0]
        for h in self.row_height:
            y_positions.append(y_positions[-1] + int(h))
        inset = max(0, self.line_width // 2)
        for r in range(self.rows):
            y0 = y_positions[r]
            y1 = y_positions[r + 1] - 1
            for c in range(self.cols):
                x0 = int(c * self.col_width)
                x1 = int(x0 + self.col_width) - 1
                draw.rectangle(
                    [x0 + inset, y0 + inset, x1 - inset, y1 - inset],
                    outline=self.line_color,
                    width=self.line_width,
                )
        for (r, c), text in self.cells.items():
            x = int(c * self.col_width)
            y = y_positions[r]
            bold = self.bold_cells.get((r, c), False)
            font_path = self.bold_font_path if bold else self.font_path
            try:
                font = ImageFont.truetype(font_path, self.font_size)
            except Exception:
                font = ImageFont.load_default()
            lines = []
            if text:
                paragraphs = str(text).split("\n")
                for p in paragraphs:
                    if p.strip() == "":
                        lines.append("")
                    else:
                        wrapped = textwrap.wrap(p, width=self.wrap_width) or [p]
                        lines.extend(wrapped)
            try:
                ascent, descent = font.getmetrics()
                line_height = ascent + descent
            except Exception:
                if lines:
                    try:
                        bbox = draw.textbbox((0, 0), lines[0], font=font)
                        line_height = bbox[3] - bbox[1]
                    except Exception:
                        _, h = font.getsize(lines[0])
                        line_height = h
                else:
                    line_height = getattr(font, "size", 12)
            cell_height = y_positions[r + 1] - y_positions[r]
            total_text_height = len(lines) * line_height
            block_top = y + (cell_height - total_text_height) / 2
            current_y = block_top
            for line in lines:
                draw.text((x + self.padding_left, current_y), line, fill=(0, 0, 0), font=font)
                current_y += line_height
        return img

# ---------- Sleepoptie image helper ----------
def create_sleepoptie_single_image(
    text,
    tekst_titel="title",
    tekst_itemnummer="1",
    canvas_height=None,
    max_chars_per_line=33,
    font_size=11,
    base_dir=None,
    num_columns=2,
):
    if not text or str(text).strip() == "":
        return None, None
    wrap_result = wrap_text(text.strip(), max_chars_per_line)
    wrapped_text = wrap_result["wrapped_text"]
    line_count = wrap_result["line_count"]
    if line_count == 1:
        calculated_height = 18
    else:
        calculated_height = max(10, (line_count * 15) + 5)
    if canvas_height is not None:
        height = max(calculated_height, canvas_height)
    else:
        height = calculated_height
    try:
        num_cols = int(num_columns) if num_columns and int(num_columns) > 0 else 2
    except Exception:
        num_cols = 2
    out_width = max(50, math.floor((450.0 / num_cols) - 15))
    if base_dir is None:
        try:
            base_dir = os.path.dirname(os.path.abspath(__file__))
        except Exception:
            base_dir = os.getcwd()
    template_path = os.path.join(base_dir, "500x500_template.png")
    if os.path.exists(template_path):
        try:
            tpl = Image.open(template_path).convert("RGB")
            tpl = tpl.resize((out_width, height), resample=Image.LANCZOS)
            img = tpl
        except Exception:
            img = Image.new("RGB", (out_width, height), "white")
    else:
        img = Image.new("RGB", (out_width, height), "white")
    draw = ImageDraw.Draw(img)
    verdana_path = os.path.join(base_dir, "verdana.ttf")
    try:
        font = ImageFont.truetype(verdana_path, font_size)
    except Exception:
        font = ImageFont.load_default()
    margin_x = 5
    margin_y = 3
    draw.multiline_text((margin_x, margin_y), wrapped_text, fill="black", font=font, spacing=4)
    filename = f"{tekst_titel}_{tekst_itemnummer}.png"
    return img, filename

# ---------- Streamlit UI & CSS ----------
st.set_page_config(page_title="Sleepoptie en Tabel Generator", layout="wide")
st.markdown(
    """
    <style>
      html, body, [data-testid="stAppViewContainer"] { background:#000; color:#fff; }
      .block-container { background: rgba(0,0,0,0.85); color: #fff; }
      footer { visibility: hidden; }
    </style>
    """,
    unsafe_allow_html=True,
)
st.info("Laatste Update: 2026-05-27 - Grootte van sleepopties bij 1 regel gefixt")
manual_filename = "Nieuwe Itemtypes Handleiding Invoer TOM.docx"
base_dir = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
manual_path = os.path.join(base_dir, manual_filename)
try:
    with open(manual_path, "rb") as f:
        manual_bytes = f.read()
    st.download_button(
        label="Klik hier om de handleiding voor invoer te downloaden!",
        data=manual_bytes,
        file_name=manual_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_manual",
    )
except FileNotFoundError:
    st.warning(f"Handleiding niet gevonden: {manual_filename}. Zet het bestand in de app-map ({base_dir}).")
except Exception as e:
    st.error(f"Kon handleiding niet laden: {e}")

st.caption("Links vul je informatie in, rechts zie je de plaatjes.")
mode = st.selectbox("Kies functie:", ["Tabel Maken", "Sleepopties Maken", "Forms Feedbacktool"], index=0)
left, right = st.columns([1, 1.2])

# ----------------- TABEL MAKEN (full functionality) -----------------
if mode == "Tabel Maken":
    with left:
        vakcode = st.text_input("Vakcode (optioneel)", value="", key="table_vakcode")
        tekst_titel = st.text_input("Titel van de tekst", value="title", key="table_titel")
        tekst_itemnummer = st.text_input("Item nummer (in selectie)", value="1", key="table_itemnr")
        table_type = st.selectbox(
            "Selecteer het Type sleepvraag",
            ("Type 1 (graphic gapmatch)", "Type 2 (graphic gapmatch categorize)"),
        )
        max_chars_per_line = st.number_input(
            "Kies het aantal karakters per regel",
            min_value=5,
            max_value=200,
            value=33,
            step=1,
            key="table_wrap_width",
        )
        with st.expander("Instellingen voor de afmetingen van de tabel", expanded=True):
            if table_type.startswith("Type 1"):
                rows = int(st.number_input("Hoeveel rijen heeft de tabel?", min_value=1, value=2, step=1, key="t1_rows"))
                cols = int(st.number_input("Hoeveel kolommen heeft de tabel?", min_value=1, value=2, step=1, key="t1_cols"))
                col_width = int(450 // cols)
                st.write(f"Met deze instellingen wordt de kolombreedte {col_width} pixels (450 // {cols})")
                heading_lines = 0
                has_heading_label = st.checkbox("Heeft de kop van de tabel een label?", value=False, key="t1_heading_has_label")
                if has_heading_label:
                    heading_lines = int(
                        st.number_input(
                            "Hoeveel regels zijn nodig voor het langste label in de kop?",
                            min_value=1,
                            value=1,
                            step=1,
                            key="heading_lines_type1",
                        )
                    )
                    st.write(f"De eerste rij van de tabel (de kop van de tabel) wordt {heading_lines * 18} pixels")
                longest_rows = int(
                    st.number_input(
                        "Hoeveel regels zijn nodig voor het langste antwoord?",
                        min_value=1,
                        value=1,
                        step=1,
                        key="t1_longest_rows",
                    )
                )
                if longest_rows == 1:
                    answer_row_height = 22
                else:
                    answer_row_height = int(longest_rows * 20)
                if heading_lines > 0:
                    first_row_height = heading_lines * 18
                    if rows == 1:
                        row_heights = [first_row_height]
                    else:
                        row_heights = [first_row_height] + [answer_row_height] * (rows - 1)
                else:
                    row_heights = [answer_row_height] * rows
            else:
                rows = 2
                cols = 2
                col_width = 225
                st.write(f"De kolombreedte voor Type 2 is altijd hetzelfde: {col_width} pixels")
                heading_lines = int(st.number_input("Hoeveel regels zijn nodig voor de kop van de tabel?", min_value=1, value=1, step=1, key="heading_rows_type2"))
                longest_rows = int(st.number_input("Hoeveel regels zijn nodig voor het langste antwoord?", min_value=1, value=1, step=1, key="longest_rows_type2"))
                answers_per_box = int(st.number_input("Hoeveel sleepopties moeten er in 1 sleepvak kunnen?", min_value=1, value=1, step=1, key="answers_per_box"))
                row1_height = int(heading_lines * 18)
                row2_height = int(longest_rows * 20 * answers_per_box)
                row_heights = [row1_height, row2_height]
        table = TableImage(rows=rows, cols=cols, row_height=row_heights, col_width=col_width, font_size=11, line_width=1, wrap_width=max_chars_per_line)
        st.subheader("Vul het benodigde tekst per cel van de tabel in")
        if table_type.startswith("Type 2"):
            bold_choice = st.checkbox("Vink dit aan als de tekst dikgedrukt moet zijn", key="table_bold_all")
        else:
            bold_choice = False
        for r in range(rows):
            cols_inputs = st.columns(cols)
            for c in range(cols):
                text_key = f"text_{r}_{c}"
                with cols_inputs[c]:
                    text = st.text_input(f"Cell {r+1},{c+1}", key=text_key)
                    if table_type.startswith("Type 1"):
                        bold_cell = st.checkbox(f"maak tekst dikgedrukt", key=f"bold_{r}_{c}")
                    else:
                        bold_cell = bold_choice
                table.set_text(r, c, text, bold=bold_cell)
    with right:
        st.subheader("Preview & Downloaden (Tabel)")
        try:
            img = table.draw()
            prefix = (safe_filename(vakcode) + "_") if (vakcode and str(vakcode).strip()) else ""
            fname_safe = prefix + (safe_filename(f"{tekst_titel}_{tekst_itemnummer}_tabel.png") or "table.png")
            st.image(img, caption=f"Gegenereerde Tabel — {fname_safe}", use_column_width=True)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            byte_im = buf.getvalue()
            st.download_button(label="Download het plaatje", data=byte_im, file_name=fname_safe, mime="image/png", key="download_table")
        except Exception as e:
            st.error(f"Kon tabel niet genereren: {e}")

# ----------------- SLEEPOPTIES MAKEN (full functionality) -----------------
elif mode == "Sleepopties Maken":
    with left:
        vakcode = st.text_input("Vakcode (optioneel)", value="", key="sleep_vakcode")
        st.header("Sleepopties genereren")
        tekst_titel = st.text_input("Titel van de tekst", value="title", key="sleep_titel")
        tekst_itemnummer = st.text_input("Item nummer (in selectie)", value="1", key="sleep_itemnr")
        num_columns = st.number_input("Hoeveel kolommen heeft de tabel?", min_value=1, value=2, step=1, key="sleep_num_columns")
        max_chars_per_line_sleep = st.number_input("Kies het aantal karakters per regel", min_value=10, value=33, key="sleep_wrap")
        st.subheader("Sleepopties (antwoordopties)")
        st.write("Laat sleepopties die je niet nodig hebt leeg.")
        letters = ["A", "B", "C", "D", "E", "F", "G", "H"]
        options = [""] * 8
        opt_cols = st.columns(2)
        for idx, L in enumerate(letters):
            with opt_cols[idx % 2]:
                options[idx] = st.text_area(f"Sleepoptie {L}", value="", height=80, key=f"opt_{L}")
    with right:
        st.subheader("Gegenereerde plaatjes")
        base_dir = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
        heights = []
        line_counts = []
        for text in options:
            if text and text.strip() != "":
                wr = wrap_text(text.strip(), max_chars_per_line_sleep)
                lc = wr["line_count"]
                if lc == 1:
                    h = 18
                else:
                    h = max(10, (lc * 15) + 5)
                heights.append(h)
                line_counts.append(lc)
        if not heights:
            st.warning("Vul op z'n minst 1 sleepoptie in.")
        else:
            max_height = max(heights)
            max_lines = max(line_counts) if line_counts else 0
            extra_pixels = max(0, max_lines - 1)
            canvas_height = max_height + extra_pixels
            generated_images = []
            prefix = (safe_filename(vakcode) + "_") if (vakcode and str(vakcode).strip()) else ""
            for idx, text in enumerate(options, start=1):
                if not text or text.strip() == "":
                    continue
                letter = chr(64 + idx)
                img, _= create_sleepoptie_single_image(
                    text,
                    tekst_titel=tekst_titel,
                    tekst_itemnummer=f"{tekst_itemnummer}_{letter}",
                    canvas_height=canvas_height,
                    max_chars_per_line=max_chars_per_line_sleep,
                    font_size=11,
                    base_dir=base_dir,
                    num_columns=num_columns,
                )
                if img is not None:
                    filename = prefix + f"{safe_filename(tekst_titel)}_{tekst_itemnummer}_{letter}.png"
                    generated_images.append((filename, img))
            if generated_images:
                st.success(f"Gegenereerde {len(generated_images)} plaatjes:")
                for i, (filename, img) in enumerate(generated_images, start=1):
                    st.image(img, caption=filename, use_column_width=False)
                    buf = io.BytesIO()
                    img.save(buf, format="PNG")
                    buf.seek(0)
                    st.download_button(label=f"Download {filename}", data=buf.getvalue(), file_name=safe_filename(filename) or filename, mime="image/png", key=f"download_{filename}_{i}")
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
                    for filename, img in generated_images:
                        img_bytes = io.BytesIO()
                        img.save(img_bytes, format="PNG")
                        img_bytes.seek(0)
                        zip_file.writestr(safe_filename(filename) or filename, img_bytes.getvalue())
                zip_buffer.seek(0)
                zip_name = prefix + f"{safe_filename(tekst_titel)}_{tekst_itemnummer}_alle_sleepopties.zip"
                st.download_button(label="Download alle plaatjes in 1 keer (.zip)", data=zip_buffer.getvalue(), file_name=safe_filename(zip_name) or "sleepopties.zip", mime="application/zip", key="download_all_zip")

# ----------------- FORMS FEEDBACKTOOL (REDESIGNED) -----------------
elif mode == "Forms Feedbacktool":
    # We need python-docx
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
        from docx.table import Table
    except Exception:
        st.error("Deze feature vereist package python-docx. Installeer met: pip install python-docx")
        st.stop()

    # Helpers for Word formatting & table borders
    def ensure_table_grid(table):
        tbl = table._tbl  # lxml element
        tblPr = None
        for child in tbl:
            if child.tag == qn("w:tblPr"):
                tblPr = child
                break
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        existing = None
        for child in tblPr:
            if child.tag == qn("w:tblBorders"):
                existing = child
                break
        if existing is not None:
            tblPr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = OxmlElement(f"w:{border_name}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "4")
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), "000000")
            borders.append(node)
        tblPr.append(borders)

    def format_para_no_spacing(para, font_family, font_size_pt, bold=False, text_override=None):
        try:
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
        except Exception:
            pass
        text = text_override if text_override is not None else (para.text or "")
        para.text = ""
        run = para.add_run(text)
        run.font.size = Pt(font_size_pt)
        run.font.bold = bold
        try:
            run.font.name = font_family
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
        except Exception:
            pass
        return run

    # helper to iterate blocks in order (paragraphs and tables)
    def iter_block_items(doc):
        for child in doc.element.body:
            tag = child.tag
            if tag.endswith("}p"):
                yield ("p", Paragraph(child, doc))
            elif tag.endswith("}tbl"):
                yield ("tbl", Table(child, doc))

    # Normalize heading for matching across docs (ignore case + collapse whitespace)
    def normalize_heading(h: str) -> str:
        if h is None:
            return ""
        h = str(h).strip()
        h = re.sub(r"\s+", " ", h)
        return h.lower()

    st.header("Forms Feedbacktool — nieuw ontwerp")
    st.write("Deze tool heeft twee secties: 'Feedbackformulieren genereren' en 'Feedbackformulieren samenvoegen'.")
    tab = st.radio("Kies sectie:", ("Feedbackformulieren genereren", "Feedbackformulieren samenvoegen"))

    # Persistent session_state initialization
    if "ff_generated" not in st.session_state:
        st.session_state["ff_generated"] = []
    if "merge_generated" not in st.session_state:
        st.session_state["merge_generated"] = []  # list of {"fname":..., "data":...}
    if "merge_ready" not in st.session_state:
        st.session_state["merge_ready"] = False

    # ---------------- Section 1: generate (Times New Roman default) ----------------
    if tab == "Feedbackformulieren genereren":
        st.subheader("1) Feedbackformulieren genereren")
        st.write("Geef hieronder op voor welke teksten je feedbackformulieren wilt aanmaken en welke onderdelen per tekst feedback moeten krijgen.")

        # Top-level inputs
        num_texts = st.number_input("Hoeveel teksten wil je invoeren?", min_value=1, max_value=50, value=1, step=1, key="ff_num_texts")
        texts = []
        for i in range(int(num_texts)):
            with st.expander(f"Tekst #{i+1}", expanded=(i == 0)):
                cg = st.text_input(f"CG (bijv. CG5 of vrij tekst) voor tekst #{i+1}", key=f"ff_cg_{i}")
                vc_type = st.selectbox(f"VC type voor tekst #{i+1}", options=["2F", "3F"], key=f"ff_vc_{i}")
                tekst_titel = st.text_input(f"Teksttitel voor tekst #{i+1}", key=f"ff_title_{i}")
                soort = st.selectbox(f"Soort tekst/items voor tekst #{i+1}", options=["Checklist", "Items", "Bezem", "Anders"], key=f"ff_type_{i}")

                feedback_points = []

                if soort == "Checklist":
                    st.markdown("Kies of je feedback wilt op:")
                    ch1 = st.checkbox("Geschiktheid Checklist", key=f"ff_check_geschik_{i}")
                    ch2 = st.checkbox("Opmerkingen", key=f"ff_check_opm_{i}")
                    if ch1:
                        feedback_points.append("Geschiktheid Checklist")
                    if ch2:
                        feedback_points.append("Opmerkingen")

                elif soort == "Items":
                    st.markdown("Kies of je feedback wilt op:")
                    it1 = st.checkbox("Algemene opmerkingen", key=f"ff_items_alg_{i}")
                    it2 = st.checkbox("Titel", key=f"ff_items_titel_{i}")
                    items_count = st.number_input("Over hoeveel items wil je feedback krijgen?", min_value=0, value=0, step=1, key=f"ff_items_count_{i}")
                    if it1:
                        feedback_points.append("Algemene opmerkingen")
                    if it2:
                        feedback_points.append("Titel")
                    for itn in range(int(items_count)):
                        feedback_points.append(f"Item {itn+1}")

                elif soort == "Bezem":
                    st.markdown("Bezem: kies of je algemene opmerkingen wilt en vink welke bezems het betreft (1 t/m 20):")
                    bz1 = st.checkbox("Algemene opmerkingen", key=f"ff_bezem_alg_{i}")
                    bezem_choices = [str(n) for n in range(1, 21)]
                    bezems_selected = st.multiselect("Welke bezems?", options=bezem_choices, key=f"ff_bezem_select_{i}")
                    if bz1:
                        feedback_points.append("Algemene opmerkingen")
                    for b in bezems_selected:
                        feedback_points.append(f"Bezem {b}")

                elif soort == "Anders":
                    st.markdown("Voer zelf op waar je feedback op wilt. Eén regel = één feedbackpunt.")
                    anders_raw = st.text_area("Vul hier de namen van de feedbackpunten (één per regel)", key=f"ff_anders_{i}", height=120)
                    if anders_raw and anders_raw.strip():
                        for ln in anders_raw.splitlines():
                            if ln.strip():
                                feedback_points.append(ln.strip())

                texts.append({
                    "cg": cg.strip() if isinstance(cg, str) else str(cg),
                    "vc_type": vc_type,
                    "titel": tekst_titel.strip() if isinstance(tekst_titel, str) else str(tekst_titel),
                    "soort": soort,
                    "feedback_points": feedback_points,
                })

        st.markdown("---")
        st.subheader("VC-instellingen")
        vc_date = st.date_input("Datum van VC", value=date.today(), key="ff_vc_date")
        vc_count = st.number_input("Aantal VC-leden", min_value=1, value=1, step=1, key="ff_vc_count")
        vc_names = []
        for vi in range(int(vc_count)):
            name = st.text_input(f"Naam VC-lid #{vi+1}", key=f"ff_vc_name_{vi}")
            vc_names.append(name.strip())

        st.markdown("---")
        # Default font must be Times New Roman
        font_family = st.selectbox("Lettertype voor Word (word-compatibel)", ["Times New Roman", "Calibri", "Arial"], index=0, key="ff_font")
        font_size_pt = st.number_input("Lettergrootte (pt) voor Word", min_value=8, max_value=18, value=11, step=1, key="ff_pt")

        generate = st.button("Genereer feedbackdocumenten", key="ff_generate")

        if generate:
            status = st.empty()
            total_points = 0
            for t in texts:
                total_points += len(t["feedback_points"])
            if total_points == 0:
                status.error("Er zijn geen feedbackpunten opgegeven. Vul voor minstens één tekst aan waarop feedback moet komen.")
            else:
                generated = []
                today_str = vc_date.isoformat()
                for vc_name in vc_names:
                    if not vc_name or vc_name.strip() == "":
                        continue
                    doc = Document()
                    for t in texts:
                        cg = t.get("cg", "")
                        titel = t.get("titel", "")
                        fp_list = t.get("feedback_points", [])
                        for fp in fp_list:
                            heading_text_parts = []
                            if cg:
                                heading_text_parts.append(cg)
                            if titel:
                                heading_text_parts.append(titel)
                            heading_main = ", ".join([p for p in heading_text_parts if p])
                            if heading_main:
                                heading_text = f"{heading_main}. {fp}" if fp else heading_main
                            else:
                                heading_text = fp or ""
                            p_head = doc.add_paragraph()
                            run_h = p_head.add_run(heading_text)
                            run_h.font.bold = True
                            run_h.font.size = Pt(font_size_pt + 1)
                            try:
                                run_h.font.name = font_family
                                run_h._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
                            except Exception:
                                pass
                            p_date = doc.add_paragraph()
                            run_d = p_date.add_run(f"FB voor VC {today_str}")
                            run_d.font.bold = True
                            run_d.font.size = Pt(font_size_pt)
                            try:
                                run_d.font.name = font_family
                                run_d._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
                            except Exception:
                                pass
                            table = doc.add_table(rows=1, cols=2)
                            ensure_table_grid(table)
                            hdr_cells = table.rows[0].cells
                            left_hdr = hdr_cells[0].paragraphs[0]
                            left_hdr.text = "VC lid"
                            format_para_no_spacing(left_hdr, font_family, font_size_pt, bold=True, text_override="VC lid")
                            right_header_text = fp if fp else "Opmerking"
                            right_hdr = hdr_cells[1].paragraphs[0]
                            right_hdr.text = right_header_text
                            format_para_no_spacing(right_hdr, font_family, font_size_pt, bold=True, text_override=right_header_text)
                            row_cells = table.add_row().cells
                            left_cell_para = row_cells[0].paragraphs[0]
                            left_cell_para.text = vc_name
                            format_para_no_spacing(left_cell_para, font_family, font_size_pt, bold=False, text_override=vc_name)
                            right_cell_para = row_cells[1].paragraphs[0]
                            right_cell_para.text = ""
                            format_para_no_spacing(right_cell_para, font_family, font_size_pt, bold=False, text_override="")
                            doc.add_page_break()
                    bio = io.BytesIO()
                    fname = f"{today_str}_VC_{safe_filename(vc_name)}_FB.docx"
                    try:
                        doc.save(bio)
                        bio.seek(0)
                        generated.append({"fname": fname, "data": bio.read()})
                    except Exception as e:
                        status.error(f"Fout bij opslaan document voor {vc_name}: {e}")
                # store persistent
                st.session_state["ff_generated"] = generated
                if not generated:
                    status.warning("Er zijn geen documenten gegenereerd (mogelijk namen van VC-leden leeg).")
                else:
                    status.success(f"✅ {len(generated)} document(en) gegenereerd en klaargezet voor download.")

        if st.session_state.get("ff_generated"):
            st.markdown("### Beschikbare gegenereerde documenten")
            for idx, item in enumerate(st.session_state["ff_generated"], start=1):
                fname = item.get("fname")
                data_bytes = item.get("data")
                if not fname or not data_bytes:
                    continue
                # stable key based on filename
                key = f"ff_dl_{safe_filename(fname)}"
                st.download_button(
                    label=fname,
                    data=data_bytes,
                    file_name=safe_filename(fname) or fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=key
                )
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
                for item in st.session_state["ff_generated"]:
                    z.writestr(safe_filename(item["fname"]) or item["fname"], item["data"])
            zip_buf.seek(0)
            try:
                chosen_date_for_zip = st.session_state.get("ff_vc_date", date.today()).isoformat()
            except Exception:
                chosen_date_for_zip = date.today().isoformat()
            zip_name = f"FB_Gebundeld_{chosen_date_for_zip}.zip"
            st.download_button(label="Download alle documenten als ZIP", data=zip_buf.getvalue(), file_name=safe_filename(zip_name) or zip_name, mime="application/zip", key="ff_dl_zip")

    # ---------------- Section 2: merging uploaded docx (per-table bundling) ----------------
    else:
        st.subheader("2) Feedbackformulieren samenvoegen")
        st.write("Upload 2 of meer .docx bestanden. Tabellen met dezelfde kop (de alinea vóór 'FB voor VC ...') worden per tabel gebundeld en in één document per CG samengevoegd, in originele volgorde.")

        uploaded = st.file_uploader("Upload .docx bestanden (meerdere mogelijk)", type=["docx"], accept_multiple_files=True)
        merge_btn = st.button("Start samenvoegen")

        # helper: parse a document into ordered list of table entries
        def parse_doc_tables(doc):
            entries = []
            last_heading_para = ""  # echte kopregel voor de tabel
            last_date_iso = None
            for block_type, block in iter_block_items(doc):
                if block_type == "p":
                    txt = (block.text or "").strip()
                    if not txt:
                        continue
                    # Detect date paragraphs like 'FB voor VC YYYY-MM-DD'
                    m_full = re.search(r"fb\s*voor\s*vc\s*(\d{4}-\d{2}-\d{2})", txt, flags=re.IGNORECASE)
                    m_iso = re.search(r"(\d{4}-\d{2}-\d{2})", txt)
                    if m_full:
                        last_date_iso = m_full.group(1)
                        continue
                    elif m_iso and txt.lower().startswith("fb"):
                        last_date_iso = m_iso.group(1)
                        continue
                    else:
                        last_heading_para = txt
                elif block_type == "tbl":
                    table = block
                    if not table.rows:
                        continue
                    hdr_cells = [c.text.strip() for c in table.rows[0].cells]
                    data_rows = []
                    for row in table.rows[1:]:
                        row_texts = [cell.text.strip() for cell in row.cells]
                        if any(cell != "" for cell in row_texts):
                            data_rows.append(row_texts)
                    heading_text = last_heading_para or "Onbekende kop"
                    entries.append({
                        "heading": heading_text,
                        "date": last_date_iso,
                        "header": hdr_cells,
                        "rows": data_rows
                    })
            return entries

        if uploaded and len(uploaded) < 2:
            st.warning("Upload minimaal 2 bestanden om samen te voegen.")

        if merge_btn:
            if not uploaded or len(uploaded) < 2:
                st.error("Je moet minstens 2 .docx bestanden uploaden.")
            else:
                status = st.empty()
                all_docs_entries = []
                dates_found = []
                filenames = [getattr(f, "name", "uploaded.docx") for f in uploaded]

                # Parse
                for f in uploaded:
                    try:
                        doc = Document(f)
                    except Exception as e:
                        status.error(f"Kon bestand niet openen: {getattr(f,'name','(naamloos)')}: {e}")
                        continue
                    entries = parse_doc_tables(doc)
                    all_docs_entries.append(entries)
                    for e in entries:
                        if e.get("date"):
                            dates_found.append(e["date"])

                if not all_docs_entries or sum(len(es) for es in all_docs_entries) == 0:
                    st.warning("Geen tabellen gevonden in de geüploade documenten.")
                else:
                    master_entries = all_docs_entries[0]
                    if not master_entries:
                        st.error("Het eerste document bevat geen tabellen; kan niet als referentie dienen.")
                        st.stop()

                    # Build canonical mapping from first document by normalized heading
                    canonical_order = []
                    canon_map = {}  # norm_heading -> {"orig": heading, "header": header, "rows": [] , "cg": cg, "order": n}
                    order_counter = 0
                    for e in master_entries:
                        h = e["heading"]
                        norm = normalize_heading(h)
                        if norm in canon_map:
                            suffix = 2
                            new_norm = f"{norm}#{suffix}"
                            while new_norm in canon_map:
                                suffix += 1
                                new_norm = f"{norm}#{suffix}"
                            norm = new_norm
                        order_counter += 1
                        m = re.match(r"^(CG\s*\d+|\d+)\b", h, flags=re.IGNORECASE)
                        if m:
                            cg_prefix = m.group(0).replace(" ", "")
                            cg_prefix = cg_prefix.upper()
                        else:
                            cg_prefix = "UNGROUPED"
                        canon_map[norm] = {
                            "orig": h,
                            "header": e.get("header", ["VC lid", "Opmerking"]),
                            "rows": list(e.get("rows", [])),
                            "cg": cg_prefix,
                            "order": order_counter,
                        }
                        canonical_order.append(norm)

                    # Function to find a matching canonical key for a heading in another doc
                    def match_canonical(norm_heading, used_set):
                        if norm_heading in canon_map and norm_heading not in used_set:
                            return norm_heading
                        base = norm_heading.split("#")[0]
                        candidates = [k for k in canonical_order if k.split("#")[0] == base and k not in used_set]
                        if candidates:
                            return candidates[0]
                        return None

                    # Merge rows from the remaining documents by heading text (normalized)
                    for doc_idx, entries in enumerate(all_docs_entries[1:], start=2):
                        used_in_this_doc = set()
                        for e in entries:
                            n = normalize_heading(e["heading"])
                            key = match_canonical(n, used_in_this_doc)
                            if key is None:
                                if len(entries) == len(canonical_order):
                                    pos = entries.index(e)
                                    if 0 <= pos < len(canonical_order):
                                        candidate = canonical_order[pos]
                                        if candidate not in used_in_this_doc:
                                            key = candidate
                            if key is None:
                                status.warning(f"Kon kop niet matchen: '{e['heading']}' in document #{doc_idx}. Deze tabel wordt overgeslagen.")
                                continue
                            used_in_this_doc.add(key)
                            canon_map[key]["rows"].extend(e.get("rows", []))
                            if e.get("date"):
                                dates_found.append(e["date"])

                    # Determine date for output
                    unique_dates = sorted(set(dates_found))
                    if len(unique_dates) == 0:
                        st.warning("Geen datums gevonden in de geüploade documenten. Kies handmatig de datum voor de gebundelde bestanden.")
                        chosen_date = st.date_input("Datum voor gebundelde bestanden", value=date.today())
                        chosen_date = chosen_date.isoformat()
                    elif len(unique_dates) == 1:
                        chosen_date = unique_dates[0]
                    else:
                        chosen_date = st.selectbox("Meerdere datums gevonden. Kies de datum voor de bestandsnamen:", options=unique_dates, index=0)

                    # Group per CG and create documents; each heading becomes its own table
                    grouped = {}
                    for key in sorted(canon_map.keys(), key=lambda k: canon_map[k]["order"]):
                        info = canon_map[key]
                        cg = info.get("cg", "UNGROUPED") or "UNGROUPED"
                        grouped.setdefault(cg, []).append(info)

                    generated = []
                    for cg_prefix, infos in grouped.items():
                        doc = Document()
                        wrote_any = False
                        for info in infos:
                            heading = info["orig"]
                            header_cells = info.get("header", ["VC lid", "Opmerking"])
                            rows = info.get("rows", [])
                            p_head = doc.add_paragraph()
                            run_h = p_head.add_run(heading)
                            run_h.font.bold = True
                            run_h.font.size = Pt(12)
                            try:
                                run_h.font.name = "Times New Roman"
                                run_h._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                            except Exception:
                                pass
                            p_date = doc.add_paragraph()
                            run_d = p_date.add_run(f"Gebundelde FB voor VC {chosen_date}")
                            run_d.font.bold = True
                            run_d.font.size = Pt(11)
                            try:
                                run_d.font.name = "Times New Roman"
                                run_d._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                            except Exception:
                                pass
                            cols_count = max(2, len(header_cells))
                            table = doc.add_table(rows=1, cols=cols_count)
                            ensure_table_grid(table)
                            for ci in range(cols_count):
                                hdr_text = header_cells[ci] if ci < len(header_cells) else ""
                                para = table.rows[0].cells[ci].paragraphs[0]
                                para.text = hdr_text
                                format_para_no_spacing(para, "Times New Roman", 11, bold=True, text_override=hdr_text)
                            for rdata in rows:
                                row_cells = table.add_row().cells
                                for ci in range(cols_count):
                                    val = rdata[ci] if ci < len(rdata) else ""
                                    para = row_cells[ci].paragraphs[0]
                                    para.text = val
                                    format_para_no_spacing(para, "Times New Roman", 11, bold=False, text_override=val)
                            doc.add_page_break()
                            wrote_any = True
                        if wrote_any:
                            fname = f"{cg_prefix}_VC{chosen_date}_FB_Gebundeld.docx"
                            bio = io.BytesIO()
                            try:
                                doc.save(bio)
                                bio.seek(0)
                                generated.append({"fname": fname, "data": bio.read()})
                            except Exception as e:
                                st.error(f"Fout bij opslaan {cg_prefix}: {e}")

                    # Persist generation as replacement so it survives reruns and downloads
                    st.session_state["merge_generated"] = generated
                    st.session_state["merge_ready"] = True

                    if not st.session_state["merge_generated"]:
                        st.warning("Er zijn geen gebundelde documenten gemaakt (mogelijk geen matchende tabellen).")
                    else:
                        st.success(f"✅ {len(st.session_state['merge_generated'])} gebundelde document(en) aangemaakt.")

        # Render download buttons persistently if merge_ready
        if st.session_state.get("merge_ready") and st.session_state.get("merge_generated"):
            st.markdown("### Beschikbare gebundelde documenten (sessie)")
            for item in st.session_state["merge_generated"]:
                fname = item.get("fname")
                data = item.get("data")
                if not fname or not data:
                    continue
                # stable key per filename ensures button remains stable across reruns
                dl_key = f"merge_dl_{safe_filename(fname)}"
                st.download_button(
                    label=fname,
                    data=data,
                    file_name=safe_filename(fname) or fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=dl_key
                )
            # ZIP button
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
                for item in st.session_state["merge_generated"]:
                    z.writestr(safe_filename(item["fname"]) or item["fname"], item["data"])
            zip_buf.seek(0)
            # try to infer chosen_date, fallback to today
            try:
                zip_date = chosen_date
            except Exception:
                zip_date = date.today().isoformat()
            zip_name = f"FB_Gebundeld_{zip_date}.zip"
            st.download_button(label="Download alle gebundelde documenten (.zip)", data=zip_buf.getvalue(), file_name=safe_filename(zip_name) or zip_name, mime="application/zip", key="merge_zip_dl")

        # Also show earlier generated files from prior runs in this session (same rendering)
        if st.session_state.get("merge_generated") and not st.session_state.get("merge_ready"):
            # this case: files exist but merge_ready False (shouldn't normally happen) — still show them
            st.markdown("### Eerder aangemaakte gebundelde documenten in deze sessie")
            for item in st.session_state["merge_generated"]:
                fname = item.get("fname")
                data = item.get("data")
                if not fname or not data:
                    continue
                dl_key = f"merge_prev_dl_{safe_filename(fname)}"
                st.download_button(
                    label=fname,
                    data=data,
                    file_name=safe_filename(fname) or fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=dl_key
                )
